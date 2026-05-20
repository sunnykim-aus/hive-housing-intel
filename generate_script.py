"""
Generates the HIVE Presentation Script as a Word document.
Run: python3 generate_script.py
Output: HIVE_Presentation_Script.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

TODAY = datetime.date.today().strftime("%-d %B %Y")

NAVY   = RGBColor(0x0F, 0x0F, 0x2A)
GOLD   = RGBColor(0xC8, 0xA0, 0x00)
DKGREY = RGBColor(0x33, 0x33, 0x33)
MGREY  = RGBColor(0x66, 0x66, 0x66)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
RED    = RGBColor(0xC0, 0x39, 0x2B)
GREEN  = RGBColor(0x1A, 0x7A, 0x40)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

doc.styles["Normal"].font.name  = "Calibri"
doc.styles["Normal"].font.size  = Pt(11)
doc.styles["Normal"].font.color.rgb = DKGREY


def spacer(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def heading(text, level=1):
    sizes = {1: 22, 2: 15, 3: 12}
    colors = {1: NAVY, 2: NAVY, 3: GOLD}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20 if level == 1 else 14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes.get(level, 11))
    r.font.color.rgb = colors.get(level, DKGREY)


def divider_bar(label, timing, color="1A1A2E"):
    """Dark section header bar with timing."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    left  = tbl.rows[0].cells[0]
    right = tbl.rows[0].cells[1]
    set_cell_bg(left,  color)
    set_cell_bg(right, color)

    left.paragraphs[0].clear()
    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = WHITE

    right.paragraphs[0].clear()
    p2 = right.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(6)
    p2.paragraph_format.space_after  = Pt(6)
    r2 = p2.add_run(timing)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0xF6, 0xC9, 0x0E)
    r2.bold = True

    spacer()


def script_para(text, space_after=10):
    """Main script body — what the speaker says."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(11.5)
    r.font.color.rgb = DKGREY
    r.font.name = "Calibri"
    return p


def stage_note(text):
    """[Stage direction] — what to click or show."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "FFF8E1")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBdr")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "C8A000")
        tcBdr.append(el)
    tcPr.append(tcBdr)

    cell.paragraphs[0].clear()
    cp = cell.add_paragraph()
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after  = Pt(4)
    cr = cp.add_run("[ " + text + " ]")
    cr.italic = True
    cr.font.size = Pt(9.5)
    cr.font.color.rgb = RGBColor(0x7A, 0x60, 0x00)
    cr.font.name = "Calibri"

    spacer()


def pause_cue(text="Pause — let them look."):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run("↳  " + text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MGREY


# ── COVER ────────────────────────────────────────────────────────────────────

spacer(3)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("🐝  HIVE")
r.bold = True
r.font.size = Pt(34)
r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(4)
r2 = p2.add_run("Housing Intelligence & Evidence")
r2.font.size = Pt(13)
r2.font.color.rgb = MGREY

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(4)
r3 = p3.add_run("Platform Walkthrough — Presentation Script")
r3.bold = True
r3.font.size = Pt(14)
r3.font.color.rgb = NAVY

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run(f"Sunny Kim  ·  Housing Data Lead  ·  {TODAY}")
r4.font.size = Pt(10)
r4.font.color.rgb = MGREY

spacer(2)

tbl = doc.add_table(rows=1, cols=1)
tbl.style = "Table Grid"
c = tbl.rows[0].cells[0]
set_cell_bg(c, "0F0F1A")
c.paragraphs[0].clear()
cp = c.add_paragraph()
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp.paragraph_format.space_before = Pt(12)
cp.paragraph_format.space_after  = Pt(12)
cr = cp.add_run(
    "Total running time: ~18–20 minutes\n"
    "Format: live platform demo — have HIVE open at localhost:8502 before you begin"
)
cr.font.size = Pt(10)
cr.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
cr.italic = True

spacer(2)

# Timing overview table
heading("Timing Overview", level=2)
sections = [
    ("Opening",                    "~2 min"),
    ("Home — the big picture",     "~2 min"),
    ("Live Dashboard",             "~2 min"),
    ("State Demand & Supply",      "~2 min"),
    ("HAFF Investment Tracker",    "~2 min"),
    ("Ask the Research",           "~3 min"),
    ("Policy Impact & Outcomes",   "~2 min"),
    ("Browse Reports & Digest",    "~1 min"),
    ("Close & Q&A invite",         "~2 min"),
    ("TOTAL",                      "~18–20 min"),
]
tbl2 = doc.add_table(rows=len(sections), cols=2)
tbl2.style = "Table Grid"
for i, (sec, t) in enumerate(sections):
    is_total = sec == "TOTAL"
    bg = "0F0F1A" if is_total else ("F5F5F5" if i % 2 == 0 else "FFFFFF")
    set_cell_bg(tbl2.rows[i].cells[0], bg)
    set_cell_bg(tbl2.rows[i].cells[1], bg)
    tbl2.rows[i].cells[0].text = sec
    tbl2.rows[i].cells[1].text = t
    tbl2.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in tbl2.rows[i].cells:
        cell.paragraphs[0].runs[0].font.size  = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE if is_total else DKGREY
        cell.paragraphs[0].runs[0].bold = is_total

doc.add_page_break()


# ── HOW TO USE THIS SCRIPT ────────────────────────────────────────────────────

heading("How to Use This Script", level=2)

script_para(
    "Read the script text aloud as written — it's phrased the way people actually speak, "
    "not formal report language. The yellow boxes are stage directions — things to click "
    "or point to on screen. You don't need to read those out."
)
script_para(
    "You don't need to memorise this word for word. Read it through twice before the session "
    "so the flow feels natural. The key is to let the platform do the talking — "
    "your job is to frame what they're looking at and why it matters."
)

spacer()
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — OPENING
# ════════════════════════════════════════════════════════════════════════════

divider_bar("OPENING", "~2 minutes", color="0F0F1A")

script_para(
    "Thanks everyone for making the time. What I want to show you today isn't a report or a "
    "presentation — it's a platform I've built specifically for our sector. I'm going to walk "
    "you through it live, so you can see exactly what it does."
)
script_para(
    "The problem it solves is one I'm sure you've all felt. Every time we need to make a case — "
    "for a funding submission, a board paper, a government meeting — we spend days hunting down "
    "the same evidence that's sitting in AHURI reports, ABS releases, AIHW data. "
    "The research exists. Getting to it is what takes the time."
)
script_para(
    "HIVE changes that. It's a web platform that puts 681 indexed housing reports, "
    "live government data, and AI-powered synthesis in one place. "
    "It's running on my machine right now — let me show you."
)

stage_note("Open browser to localhost:8502 — Home page should be visible")

spacer()
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — HOME
# ════════════════════════════════════════════════════════════════════════════

divider_bar("HOME  —  The Big Picture", "~2 minutes", color="0F0F1A")

stage_note("Point to the hero banner at the top of the Home page")

script_para(
    "This is the Home page. The first thing you'll notice at the top is four live numbers — "
    "and these aren't numbers I've typed in. They're pulled automatically from ABS and AIHW "
    "every month."
)

stage_note("Point to the four stat blocks — run rate, shortfall, unmet requests, success rate")

script_para(
    "Australia is currently building around 160,000 dwellings a year. "
    "The National Housing Accord target is 240,000. "
    "That's an 80,000-dwelling shortfall every single year."
)
pause_cue()
script_para(
    "At the same time — last year, over 100,000 people went to a homelessness service "
    "and left without housing. And only 27% of people who needed housing actually received it. "
    "That's the sector context we're operating in, right there on one screen."
)

stage_note("Scroll down to the pipeline status card")

script_para(
    "Below that you can see the data pipeline status — right now we have 681 reports indexed "
    "and over 150,000 searchable chunks ready to go. I'll come back to what that means "
    "when we get to the research search."
)

stage_note("Scroll down briefly to show the 'Who Uses HIVE' role cards — don't linger")

script_para(
    "And at the bottom of the Home page there are role guides — "
    "CEO, Policy Manager, Development, Grants — each one tells you which pages matter most "
    "for that role. But let's go through the main ones now."
)

spacer()
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — LIVE DASHBOARD
# ════════════════════════════════════════════════════════════════════════════

divider_bar("LIVE DASHBOARD", "~2 minutes", color="0F0F1A")

stage_note("Click 'Live Dashboard' in the top nav bar")

script_para(
    "The Live Dashboard is where you go when you need current numbers. "
    "Everything here is sourced directly from the ABS and AIHW — "
    "no manual updates, no stale spreadsheets."
)

stage_note("Point to the building approvals trend chart")

script_para(
    "This top chart shows monthly building approvals going back 10 years — "
    "total dwellings, houses, and other types like apartments. "
    "You can see the COVID spike in 2021 and the drop-off since. "
    "That decline is what's driving the shortfall we just saw on the Home page."
)

stage_note("Scroll down to the SHS homelessness funnel")

script_para(
    "This funnel is what I find the most confronting. It starts with the total number of "
    "people who came to a Specialist Homelessness Service last year — then shows how many "
    "needed housing, how many were assigned housing, and how many actually achieved "
    "stable housing by the end. Watch that drop-off. That's the system working at capacity "
    "and still not meeting demand."
)
pause_cue()

stage_note("Scroll down to the state waitlist comparison")

script_para(
    "And down here are the state waitlists — approved applicants by state, year-on-year change, "
    "and how many years at the current build rate before the list clears. "
    "For context, NSW is sitting at over 30 years."
)

spacer()
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — STATE DEMAND & SUPPLY
# ════════════════════════════════════════════════════════════════════════════

divider_bar("STATE DEMAND & SUPPLY", "~2 minutes", color="0F0F1A")

stage_note("Click 'Demand & Supply' in the nav bar")

script_para(
    "This page is particularly useful for development teams and anyone building a "
    "state-specific investment case. Let me show you what I mean."
)

stage_note("Select a state from the dropdown — e.g. NSW or VIC")

script_para(
    "I've selected New South Wales. What you're looking at now is who is actually "
    "on the waitlist — by household type. Single person households, couples, families. "
    "And on the right, what's being built and whether the bedroom mix matches that need."
)

stage_note("Point to the demand vs supply mismatch chart")

script_para(
    "This is the mismatch chart. The waitlist skews heavily toward single-person households — "
    "but the approvals pipeline is building mostly two and three bedroom dwellings. "
    "Supply is going up, but it's not the supply the waitlist actually needs. "
    "That's a structural problem that no amount of general construction will fix."
)

stage_note("Scroll to show the 20-year waitlist trend")

script_para(
    "And here's 20 years of the waitlist. You can see it's not a recent spike — "
    "this is decades of accumulated under-supply. That's the long-term story in one line."
)

stage_note("Scroll to the state spotlight table at the bottom — point to all 5 states")

script_para(
    "The table at the bottom pulls all five states together — waitlist size, year-on-year change, "
    "approvals, social housing built, and years to clear. "
    "If you ever need to brief across states, this is the table."
)

spacer()
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — HAFF
# ════════════════════════════════════════════════════════════════════════════

divider_bar("HAFF INVESTMENT TRACKER", "~2 minutes", color="0F0F1A")

stage_note("Click 'HAFF' in the nav bar")

script_para(
    "The Housing Australia Future Fund is the biggest federal housing investment "
    "in a generation — $10 billion, targeting 30,000 homes over five years. "
    "This page tracks exactly where that investment is going."
)

stage_note("Point to the KPI row at the top — total homes, grants, progress bar")

script_para(
    "Across the top you can see: total grants committed, homes announced across all three rounds, "
    "and the progress toward the 30,000 target. "
    "We're at about 60% of that target with grants committed — "
    "but committed and delivered are two different things, which is why the delivery pipeline "
    "at the bottom of each round is worth watching."
)

stage_note("Click on 'Round 1 — March 2024' tab")

script_para(
    "Each round has its own tab. For Round 1 you can see the state breakdown — "
    "how many homes per state, how many are social vs affordable, "
    "the grant per home, and which community housing providers are the delivery partners."
)

stage_note("Scroll to show the sector and bedroom mix charts")

script_para(
    "These two charts are the ones I find most useful in a pitch context. "
    "The sector chart shows the target populations — family violence, Indigenous, "
    "older persons, key workers — so you can make the case that HAFF isn't just "
    "housing numbers, it's targeted housing for people with the greatest need. "
    "And the bedroom mix shows whether what's being built actually matches who's waiting."
)

spacer()
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — ASK THE RESEARCH
# ════════════════════════════════════════════════════════════════════════════

divider_bar("ASK THE RESEARCH", "~3 minutes  ★ This is the centrepiece", color="0F0F1A")

stage_note("Click 'Ask Research' in the nav bar")

script_para(
    "This is the part of the platform I'm most excited about — and the one that will probably "
    "save you the most time."
)
script_para(
    "You're looking at a search box. But this isn't a keyword search. "
    "HIVE has read and indexed 681 reports — every AHURI final report, every AIHW annual report, "
    "15 years of Treasury budget papers, Productivity Commission inquiries. "
    "All of it is sitting in a vector database, which means it understands meaning, "
    "not just words."
)

stage_note("Type a question into the search box — use: 'What does the research say about social housing supply barriers in Australia?'")

script_para(
    "I'm going to type a question I'd genuinely need answered before a policy submission. "
    "Let me show you what comes back."
)

stage_note("Click 'Search & Analyse' — let it run, point to the spinner")
pause_cue("Give it 15–20 seconds. Keep talking while it runs.")

script_para(
    "What's happening right now — HIVE is finding the most relevant passages across those "
    "681 reports, and sending them to Claude to write a synthesised answer. "
    "Not a list of links. An actual answer, with citations."
)

stage_note("Point to the synthesised answer when it appears")

script_para(
    "There it is. A paragraph-by-paragraph synthesis of what the research says, "
    "with the specific reports it's drawing from. "
    "Scroll down and you can see each source — the title, the agency, the year, "
    "and the exact passage it used."
)

stage_note("Point to the 'Export to Word' button")

script_para(
    "And then — one click. Export to Word. "
    "You get a formatted document with the analysis and all the citations, "
    "ready to paste into a submission or board paper. "
    "What would normally take a researcher half a day takes about 30 seconds."
)
pause_cue("Let that land.")

spacer()
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — POLICY IMPACT & OUTCOME LEDGER
# ════════════════════════════════════════════════════════════════════════════

divider_bar("POLICY IMPACT  &  OUTCOME LEDGER", "~2 minutes", color="0F0F1A")

stage_note("Click 'Policy Impact' in the nav bar")

script_para(
    "The next two pages are about accountability — what the government has committed "
    "to the sector and what it's actually delivered."
)
script_para(
    "Policy Impact lets you select any major program — HAFF, NRAS, HomeBuilder, "
    "Nation Building — and get an evidence-based assessment of what the research says "
    "about how it performed. Not the government's own evaluation. Independent research."
)

stage_note("Select a program from the dropdown — e.g. 'National Rental Affordability Scheme (NRAS)' — click 'Analyse Impact'")

script_para(
    "I'll run NRAS as an example — it's the most analysed program in the AHURI archive "
    "so you'll see a rich result."
)

stage_note("While it runs, navigate to Outcome Ledger")

stage_note("Click 'Outcomes' in the nav bar — point to the program list on the left")

script_para(
    "The Outcome Ledger is a different view of the same accountability question. "
    "It tracks specific commitments — the target, the actual result, and a status. "
    "On track, met, missed, or ongoing. "
    "This is the page you'd use in a board paper or an investor briefing "
    "when you need to show the track record of public investment in housing."
)

stage_note("Click on one program — point to the targets vs actuals grid")

script_para(
    "You can see for each program: what was promised, what was delivered, "
    "and where the gaps are. The research extraction button at the bottom will "
    "pull additional commentary from the indexed reports if you want deeper analysis."
)

spacer()
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 8 — BROWSE REPORTS & WEEKLY DIGEST
# ════════════════════════════════════════════════════════════════════════════

divider_bar("BROWSE REPORTS  &  WEEKLY DIGEST", "~1 minute", color="0F0F1A")

stage_note("Click 'Reports' in the nav bar — briefly show the table")

script_para(
    "Browse Reports is the full library. You can filter by source, year, and report type — "
    "and open or download any document. Every report in the index is listed here. "
    "It's useful when you know what you're looking for and just want to go straight to the source."
)

stage_note("Click 'Digest' in the nav bar")

script_para(
    "The Weekly Digest is exactly what it sounds like. "
    "Click Generate and HIVE produces a briefing combining the latest live data "
    "with any research published that week — formatted and ready to send. "
    "If you're sending a weekly update to a board or leadership team, "
    "this is how you do it without spending two hours pulling numbers together."
)

spacer()
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 9 — CLOSE
# ════════════════════════════════════════════════════════════════════════════

divider_bar("CLOSE  &  Q&A", "~2 minutes", color="0F0F1A")

stage_note("Return to Home page — let the platform sit on screen")

script_para(
    "So that's HIVE. To summarise what it does in three lines:"
)
script_para(
    "It gives you instant access to 15 years of housing research — synthesised and cited, "
    "not just a list of links."
)
script_para(
    "It keeps you current with live government data — building approvals, homelessness statistics, "
    "state waitlists — updated automatically."
)
script_para(
    "And it turns hours of evidence-gathering into minutes — with a Word export ready to go "
    "straight into a submission, a board paper, or a funding pitch."
)
pause_cue()
script_para(
    "Everything runs locally — your data doesn't go anywhere except the Claude API call "
    "when you ask for a synthesis. The pipeline runs once and stays indexed. "
    "No subscription, no ongoing cost beyond the API usage."
)
script_para(
    "I built this because I kept doing this work the hard way and knew there had to be "
    "a better way to do it. I'd love to hear your reaction — "
    "and particularly what you'd want to see added or changed to make it useful for your work."
)
script_para(
    "Happy to take questions."
)

spacer(2)

# Notes box
tbl = doc.add_table(rows=1, cols=1)
tbl.style = "Table Grid"
c = tbl.rows[0].cells[0]
set_cell_bg(c, "F5F5F5")
tc = c._tc
tcPr = tc.get_or_add_tcPr()
tcBdr = OxmlElement("w:tcBdr")
for side in ["top", "left", "bottom", "right"]:
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), "4")
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), "1A1A2E")
    tcBdr.append(el)
tcPr.append(tcBdr)

c.paragraphs[0].clear()
np_ = c.add_paragraph()
np_.paragraph_format.space_before = Pt(8)
np_.paragraph_format.space_after  = Pt(4)
nr = np_.add_run("Likely questions to prepare for")
nr.bold = True
nr.font.size = Pt(10)
nr.font.color.rgb = NAVY

for q in [
    "How much does the API cost to run? — A typical research query costs $0.01–0.03. The weekly digest ~$0.05. "
    "Most features are free (no API call).",
    "Can other people in the team access it? — Currently runs locally. Could be hosted on a server for shared access "
    "— that's a conversation about infrastructure.",
    "How often does the data update? — ABS live data refreshes monthly automatically. Indexed reports: re-run "
    "the pipeline to pick up new publications (takes 15–30 min).",
    "Is the data accurate? — Every figure links to its primary source. The AI synthesis can occasionally "
    "misinterpret a source — always verify before formal use.",
    "Can it cover state programs too? — Currently it indexes federal and major state publications. "
    "State-specific reports can be added by placing PDFs in the data folder and re-running the pipeline.",
]:
    qp = c.add_paragraph()
    qp.paragraph_format.left_indent = Cm(0.3)
    qp.paragraph_format.space_after = Pt(4)
    qr = qp.add_run("→  " + q)
    qr.font.size = Pt(9.5)
    qr.font.color.rgb = DKGREY

c.add_paragraph().paragraph_format.space_after = Pt(4)

spacer(2)

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run(f"HIVE  ·  Presentation Script  ·  {TODAY}  ·  Sunny Kim")
r.font.size  = Pt(8)
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


output_path = "HIVE_Presentation_Script.docx"
doc.save(output_path)
print(f"✅  Saved: {output_path}")
