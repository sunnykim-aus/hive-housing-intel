"""
Generates the HIVE User Guide as a Word document.
Run: python3 generate_guide.py
Output: HIVE_User_Guide.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

TODAY = datetime.date.today().strftime("%-d %B %Y")
YEAR = datetime.date.today().year

# ── Colour palette ──────────────────────────────────────────────────────────
GOLD    = RGBColor(0xF6, 0xC9, 0x0E)
NAVY    = RGBColor(0x0F, 0x0F, 0x1A)
MID     = RGBColor(0x1A, 0x1A, 0x2E)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
LGREY   = RGBColor(0xCC, 0xCC, 0xCC)
DKGREY  = RGBColor(0x44, 0x44, 0x44)
RED     = RGBColor(0xE7, 0x4C, 0x3C)
BLUE    = RGBColor(0x34, 0x98, 0xDB)
GREEN   = RGBColor(0x27, 0xAE, 0x60)


def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_para(doc, text, style="Normal", bold=False, italic=False,
             color=None, size=None, space_before=0, space_after=6,
             alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(style=style)
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p


def add_heading(doc, text, level=1):
    heading_sizes = {1: 22, 2: 16, 3: 13, 4: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(heading_sizes.get(level, 11))
    run.font.color.rgb = NAVY if level == 1 else DKGREY
    return p


def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_page_intro_table(doc, page_name: str, tagline: str, who: str, what_you_get: str):
    """Dark summary card for each page section."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "1A1A2E")
    cell.paragraphs[0].clear()

    def _add(text, bold=False, clr=WHITE, sz=10):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(sz)
        r.font.color.rgb = clr

    _add(f"🔹  {page_name}", bold=True, clr=GOLD, sz=12)
    _add(tagline, clr=LGREY, sz=10)
    _add("")
    _add("Who uses it:  " + who, clr=LGREY)
    _add("What you get: " + what_you_get, clr=LGREY)

    doc.add_paragraph()  # spacer


def add_tip(doc, tip_text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run("💡  Tip: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0xF3, 0x9C, 0x12)
    r.font.size = Pt(10)
    r2 = p.add_run(tip_text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DKGREY


# ── Build document ───────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)
doc.styles["Normal"].font.color.rgb = DKGREY


# ── COVER PAGE ──────────────────────────────────────────────────────────────

add_para(doc, "", space_before=20)

add_para(doc, "🐝  HIVE", bold=True, size=36, color=NAVY,
         space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_para(doc, "Housing Intelligence & Evidence", bold=False, size=14,
         color=DKGREY, space_after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_para(doc, "User Guide & Platform Briefing", bold=True, size=16,
         color=NAVY, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_para(doc, f"Version 1.0  ·  {TODAY}", bold=False, size=10,
         color=DKGREY, space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# Tagline box
tbl = doc.add_table(rows=1, cols=1)
tbl.style = "Table Grid"
c = tbl.rows[0].cells[0]
set_cell_bg(c, "0F0F1A")
c.paragraphs[0].clear()
p = c.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after  = Pt(12)
r = p.add_run(
    "15 years of indexed research, live government data,\n"
    "and AI-powered synthesis — built for the people working inside the sector."
)
r.font.size  = Pt(11)
r.font.color.rgb = LGREY
r.italic = True

doc.add_paragraph()

add_para(doc, "Prepared by Sunny Kim  ·  Housing Data Lead  ·  Community Housing Professional, Australia",
         size=9, color=DKGREY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "linkedin.com/in/sunny-kim-58a780100",
         size=9, color=DKGREY, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()


# ── TABLE OF CONTENTS (manual) ───────────────────────────────────────────────

add_heading(doc, "Contents", level=1)

toc_items = [
    ("1",  "Executive Summary",                 "3"),
    ("2",  "What Is HIVE?",                     "3"),
    ("3",  "Getting Started",                   "4"),
    ("4",  "Navigating the Platform",           "5"),
    ("5",  "Page-by-Page Guide",                "5"),
    ("",   "5.1   Home",                        "5"),
    ("",   "5.2   Live Dashboard",              "6"),
    ("",   "5.3   State Demand & Supply",       "7"),
    ("",   "5.4   HAFF Investment Tracker",     "8"),
    ("",   "5.5   Ask the Research",            "9"),
    ("",   "5.6   Policy Impact",               "10"),
    ("",   "5.7   Outcome Ledger",              "11"),
    ("",   "5.8   Policy Timeline",             "12"),
    ("",   "5.9   Browse Reports",              "12"),
    ("",   "5.10  Weekly Digest",               "13"),
    ("6",  "Who Uses HIVE — Role Guide",        "14"),
    ("7",  "Data Sources",                      "15"),
    ("8",  "Tips & Frequently Asked Questions", "16"),
    ("9",  "Glossary",                          "17"),
]

tbl = doc.add_table(rows=len(toc_items), cols=2)
tbl.style = "Table Grid"
for i, (num, label, page_num) in enumerate(toc_items):
    row = tbl.rows[i]
    row.cells[0].text = f"  {num}  {label}" if num else f"       {label}"
    row.cells[1].text = page_num
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = DKGREY
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(2)

doc.add_page_break()


# ── 1  EXECUTIVE SUMMARY ─────────────────────────────────────────────────────

add_heading(doc, "1.  Executive Summary", level=1)

add_para(doc, (
    "HIVE (Housing Intelligence & Evidence) is a web-based intelligence platform designed "
    "for professionals working in or with the Australian community housing sector. "
    "It addresses a practical problem: critical housing evidence exists in hundreds of "
    "government and research publications, but finding, synthesising, and acting on that "
    "evidence currently takes days of manual effort."
), space_after=8)

add_para(doc, (
    "HIVE solves this by combining three capabilities in a single platform:"
), space_after=4)

add_bullet(doc, "A searchable, AI-indexed library of 681+ research reports from AHURI, AIHW, "
           "ABS, Treasury, Housing Australia, the Productivity Commission, DSS, and Power Housing.")
add_bullet(doc, "Live government data — ABS building approvals, AIHW homelessness statistics, "
           "state housing waitlists — updated automatically and visualised in context.")
add_bullet(doc, "Claude AI synthesis — ask any housing question in plain language and receive "
           "a cited, evidence-based answer in seconds, exportable to Word.")

add_para(doc, (
    "HIVE is not a general research tool. It is purpose-built for the community housing sector: "
    "every data source, every chart, every AI prompt is calibrated to the questions that CEOs, "
    "policy managers, development teams, and grant officers actually need answered."
), space_before=8, space_after=8)


# ── 2  WHAT IS HIVE ──────────────────────────────────────────────────────────

add_heading(doc, "2.  What Is HIVE?", level=1)

add_para(doc, (
    "HIVE stands for Housing Intelligence & Evidence. The platform is a Streamlit web application "
    "running locally on your machine at localhost:8502. All data processing happens on your device — "
    "the only external call is to the Claude AI API when you request an AI synthesis or insight."
), space_after=8)

add_heading(doc, "Core capabilities", level=3)

rows = [
    ("Research synthesis",       "Search 681+ indexed reports and receive a synthesised, cited answer. "
                                  "Export directly to Word for submissions, board papers, or grant applications."),
    ("Live market intelligence",  "Real-time ABS building approvals, AIHW homelessness data, and state "
                                  "housing waitlist trends — automatically refreshed monthly."),
    ("Policy accountability",     "Track every major federal housing investment against what it promised. "
                                  "Targets vs actuals for HAFF, NRAS, Nation Building, NHFIC, and more."),
    ("State demand & supply",     "State-by-state analysis: who is on the waitlist, what household type, "
                                  "what is being built, and the structural mismatch between supply and need."),
    ("Evidence export",           "Export any analysis to a formatted Word document with citations "
                                  "— ready for board papers, investor briefings, or government submissions."),
    ("Weekly digest",             "An AI-generated weekly briefing combining live data and newly indexed "
                                  "research — keeps leadership teams current without manual compilation."),
]

tbl = doc.add_table(rows=len(rows)+1, cols=2)
tbl.style = "Table Grid"

# Header
hdr_cells = tbl.rows[0].cells
for cell, text in zip(hdr_cells, ["Capability", "Description"]):
    set_cell_bg(cell, "0F0F1A")
    cell.text = text
    cell.paragraphs[0].runs[0].font.color.rgb = GOLD
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

for i, (cap, desc) in enumerate(rows, 1):
    row = tbl.rows[i]
    if i % 2 == 0:
        set_cell_bg(row.cells[0], "F5F5F5")
        set_cell_bg(row.cells[1], "F5F5F5")
    row.cells[0].text = cap
    row.cells[1].text = desc
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].paragraphs[0].runs[0].font.color.rgb = DKGREY

doc.add_paragraph()

doc.add_page_break()


# ── 3  GETTING STARTED ───────────────────────────────────────────────────────

add_heading(doc, "3.  Getting Started", level=1)

add_heading(doc, "3.1  Prerequisites", level=2)

add_para(doc, "Before launching HIVE, ensure the following are in place:", space_after=4)
add_bullet(doc, "Python 3.11+ installed")
add_bullet(doc, "All dependencies installed (pip install -r requirements.txt in the housing-intel directory)")
add_bullet(doc, "An Anthropic API key stored in the .env file as ANTHROPIC_API_KEY=sk-ant-...")
add_bullet(doc, "At least 2 GB of free disk space for the report index")

add_heading(doc, "3.2  Launching the platform", level=2)

add_para(doc, "Open a terminal, navigate to the housing-intel directory, and run:", space_after=4)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
p.paragraph_format.space_after = Pt(8)
r = p.add_run("streamlit run app.py --server.port 8502")
r.font.name = "Courier New"
r.font.size  = Pt(10)
r.font.color.rgb = RGBColor(0x1A, 0x80, 0x40)

add_para(doc, (
    "The platform will open automatically in your default browser at http://localhost:8502. "
    "Keep the terminal window open while using HIVE — closing it stops the server."
), space_after=8)

add_heading(doc, "3.3  First-time setup — running the data pipeline", level=2)

add_para(doc, (
    "On first launch, HIVE will display a pipeline status card showing 0 reports indexed. "
    "You must run the pipeline once before the AI search and synthesis features will work."
), space_after=6)

add_para(doc, "To run the pipeline:", bold=True, size=10, space_after=4)
add_bullet(doc, "Click the 'Run All Sources — Index Reports' button on the Home page.")
add_bullet(doc, "The pipeline will crawl all 8 data sources and index every document it finds.")
add_bullet(doc, "This takes approximately 15–30 minutes on first run, depending on your connection speed.")
add_bullet(doc, "Once complete, the status card will show the number of reports indexed and 'Live & Indexed' status.")
add_bullet(doc, "The pipeline does not need to be re-run manually — it detects when data is already indexed "
           "and will not consume API credits on subsequent launches.")

add_tip(doc, "The Run button only appears when the index is empty. Once indexed, HIVE will not show a "
         "re-run option — this prevents unnecessary API and compute costs.")

add_heading(doc, "3.4  API key costs", level=2)

add_para(doc, (
    "HIVE uses the Claude API (claude-sonnet-4-6) only when: (a) you click 'Search & Analyse' "
    "on Ask the Research, (b) you click 'Analyse Impact' on Policy Impact, (c) you click "
    "'Get AI Insight' on any page, or (d) you generate the Weekly Digest. "
    "All other features — charts, tables, live data — run without any API usage."
), space_after=6)

add_para(doc, (
    "AI insights are cached in your browser session. Within a single session, the same "
    "insight will not re-trigger an API call even if you navigate away and return."
), space_after=8)

doc.add_page_break()


# ── 4  NAVIGATION ────────────────────────────────────────────────────────────

add_heading(doc, "4.  Navigating the Platform", level=1)

add_para(doc, (
    "HIVE uses a horizontal pill navigation bar at the top of every page. "
    "Click any label to switch pages instantly — no page reload required."
), space_after=6)

nav_items = [
    ("Home",           "Platform overview, live crisis numbers, data pipeline status, quick search"),
    ("Live Dashboard", "ABS building approvals, AIHW homelessness funnel, state waitlist trends"),
    ("Demand & Supply","State-by-state demand and supply analysis with household type breakdown"),
    ("HAFF",           "Housing Australia Future Fund — round-by-round project and grant breakdown"),
    ("Ask Research",   "AI-powered search across 681+ indexed reports — plain language questions"),
    ("Policy Impact",  "Evidence-based impact assessment for any major housing program"),
    ("Outcomes",       "Policy Outcome Ledger — investment promises vs delivered results"),
    ("Timeline",       "Visual timeline of major Australian housing policies since 1990"),
    ("Reports",        "Browse, filter, and download the full indexed report library"),
    ("Digest",         "AI-generated weekly sector briefing with live data and new research"),
]

tbl = doc.add_table(rows=len(nav_items)+1, cols=2)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
for cell, txt in zip(hdr, ["Nav label", "What it does"]):
    set_cell_bg(cell, "0F0F1A")
    cell.text = txt
    cell.paragraphs[0].runs[0].font.color.rgb = GOLD
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

for i, (label, desc) in enumerate(nav_items, 1):
    if i % 2 == 0:
        set_cell_bg(tbl.rows[i].cells[0], "F5F5F5")
        set_cell_bg(tbl.rows[i].cells[1], "F5F5F5")
    tbl.rows[i].cells[0].text = label
    tbl.rows[i].cells[1].text = desc
    for cell in tbl.rows[i].cells:
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = DKGREY
    tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
doc.add_page_break()


# ── 5  PAGE BY PAGE ──────────────────────────────────────────────────────────

add_heading(doc, "5.  Page-by-Page Guide", level=1)


# 5.1 Home
add_heading(doc, "5.1  Home", level=2)
add_page_intro_table(doc,
    page_name   = "Home",
    tagline     = "The entry point — live crisis context, pipeline status, and a quick research search.",
    who         = "Everyone. Start here every session.",
    what_you_get= "National housing numbers, data pipeline health, quick search, and a summary of all platform capabilities.",
)

add_heading(doc, "What's on this page", level=3)
add_bullet(doc, "Hero banner — the HIVE title card with 681 reports indexed, data source badges, and the current date.")
add_bullet(doc, "Pipeline status card — shows how many reports are indexed and the number of searchable vector chunks. "
           "Run button only appears when the index is empty.")
add_bullet(doc, "The Housing Crisis in Numbers — four live metrics: annual dwelling approvals, supply shortfall vs the "
           "National Housing Accord target, unmet SHS requests, and housing success rate. All sourced from ABS and AIHW.")
add_bullet(doc, "Evidence reference table — the data source, publication, and update frequency behind every number.")
add_bullet(doc, "What HIVE Does — six capability cards summarising each module.")
add_bullet(doc, "Who Uses HIVE — role cards for CEO, Policy Manager, Development Manager, Grants Officer, "
           "Impact Investor, and Government Stakeholder — each with recommended starting pages.")
add_bullet(doc, "Quick search bar — ask any housing question without leaving the Home page.")
add_bullet(doc, "Data sources section — descriptions of all 8 indexed sources with report and chunk counts.")

add_tip(doc, "The four live numbers update automatically each month when new ABS or AIHW data is published. "
         "If you need to force a refresh, delete the relevant file in data/live_cache/ and reload.")


# 5.2 Live Dashboard
add_heading(doc, "5.2  Live Dashboard", level=2)
add_page_intro_table(doc,
    page_name   = "Live Dashboard",
    tagline     = "Real-time national housing supply and demand data — updated monthly from ABS and AIHW.",
    who         = "CEOs, policy managers, anyone needing current sector numbers for briefings or submissions.",
    what_you_get= "Building approvals trend, SHS homelessness funnel, state waitlist comparison, and AI insights.",
)

add_heading(doc, "What's on this page", level=3)
add_bullet(doc, "National supply metrics — current annual run rate, 3-month average, year-on-year change, and gap to the "
           "240,000 National Accord target.")
add_bullet(doc, "Building approvals trend chart — monthly total, houses, and other dwellings from ABS Cat. 8731.0. "
           "Toggle between total, houses only, or other dwellings.")
add_bullet(doc, "Specialist Homelessness Services (SHS) funnel — AIHW data showing total clients, those needing housing, "
           "those assigned housing, and those who achieved housing stability. The funnel shows the attrition at each step.")
add_bullet(doc, "Homelessness trend by support type — bar chart of SHS clients by service type over time.")
add_bullet(doc, "State waitlist comparison — approved applicants by state, year-on-year change, and years-to-clear estimate.")
add_bullet(doc, "AI Insight buttons — click 'Get AI Insight' under any chart for a 2-sentence Claude analysis of what the data means.")

add_tip(doc, "The SHS data is sourced from the AIHW annual report (2023–24 is the most recent). "
         "The funnel numbers are not directly comparable year-on-year because the AIHW revises methodology.")


# 5.3 State Demand & Supply
add_heading(doc, "5.3  State Demand & Supply", level=2)
add_page_intro_table(doc,
    page_name   = "State Demand & Supply",
    tagline     = "Deep-dive analysis by state — who is on the waitlist, what is being built, and the structural mismatch.",
    who         = "Development managers, policy teams, anyone building a state-specific investment case.",
    what_you_get= "State selector, waitlist demographics, dwelling type mismatch, 20-year waitlist trend, and state spotlight table.",
)

add_heading(doc, "What's on this page", level=3)
add_bullet(doc, "State selector — choose NSW, VIC, QLD, WA, or SA for a state-specific deep dive.")
add_bullet(doc, "Waitlist demographics — household types on the waitlist (single person, couple, family) vs "
           "what HAFF and building approvals are actually delivering by dwelling type.")
add_bullet(doc, "Demand vs supply mismatch chart — side-by-side comparison of waitlist need and supply output.")
add_bullet(doc, "20-year waitlist trend — how the approved applicant count has changed since 2004.")
add_bullet(doc, "State spotlight table — all 5 states in one view: waitlist total, YoY change, total approvals, "
           "social/affordable built, % accessible, estimated years to clear, and social housing stock.")
add_bullet(doc, "AI Insight — state-specific Claude assessment of the demand-supply situation.")

add_tip(doc, "The 'Years to clear waitlist' estimate assumes the current annual social housing completion rate continues. "
         "It is a planning indicator, not a prediction.")


# 5.4 HAFF
add_heading(doc, "5.4  HAFF Investment Tracker", level=2)
add_page_intro_table(doc,
    page_name   = "HAFF Investment Tracker",
    tagline     = "Housing Australia Future Fund — round-by-round breakdown of projects, grants, and dwelling types.",
    who         = "Development managers, CEOs, policy teams, investors assessing the HAFF pipeline.",
    what_you_get= "Fund KPIs, progress to 30,000 target, state breakdown charts, sector mix, bedroom mix, and delivery pipeline.",
)

add_heading(doc, "What's on this page", level=3)
add_bullet(doc, "Fund overview — $10B fund size, total grants committed across rounds, homes announced vs 30,000 target, "
           "split between social (20,000 target) and affordable (10,000 target).")
add_bullet(doc, "Progress bar — visual progress toward the 30,000 home 5-year target with remaining allocation.")
add_bullet(doc, "Round tabs — separate tabs for Round 1 (March 2024), Round 2 (October 2024), Round 3 (March 2025), "
           "and an All Rounds Combined view.")
add_bullet(doc, "Per-round detail — context notes, status, grants total, projects, CHP delivery partners, average grant per home, "
           "completion target.")
add_bullet(doc, "State breakdown charts — homes by state (stacked social/affordable bar), grant allocation by state (horizontal bar).")
add_bullet(doc, "Sector and dwelling type mix — target populations (family violence, Indigenous, older persons, key workers, general) "
           "and dwelling type breakdown (apartments, townhouses, detached, supported).")
add_bullet(doc, "Bedroom mix — 1-bed/studio, 2-bed, 3-bed, 4+-bed counts vs waitlist household need.")
add_bullet(doc, "Delivery pipeline — milestone status for each round with completion target dates.")
add_bullet(doc, "AI Insight — Claude assessment of delivery risk relative to the 30,000 home target.")

add_tip(doc, "Source: Housing Australia media releases, Senate Estimates, and Budget Papers 2023–24 to 2025–26. "
         "Verify against primary Housing Australia sources before formal submission use.")


# 5.5 Ask the Research
add_heading(doc, "5.5  Ask the Research", level=2)
add_page_intro_table(doc,
    page_name   = "Ask the Research",
    tagline     = "Search 15 years of housing research in plain language and get a synthesised, cited answer.",
    who         = "Policy teams, grants officers, executives — anyone who needs evidence fast.",
    what_you_get= "A cited AI synthesis from up to 25 source documents, plus individual source excerpts, exportable to Word.",
)

add_heading(doc, "How to use it", level=3)
add_bullet(doc, "Type your question in plain language — no need for keywords or boolean operators.")
add_bullet(doc, "Example questions: 'What does the research say about social housing waitlist trends in Victoria?', "
           "'How effective has NRAS been at increasing affordable rental supply?', "
           "'What are the barriers to housing supply in Australian cities?'")
add_bullet(doc, "Use the Search Filters expander to narrow by year range (2005–2025) and number of sources to retrieve (5–25).")
add_bullet(doc, "Click 'Search & Analyse'. HIVE will retrieve the most semantically relevant document chunks from the index "
           "and send them to Claude to produce a synthesised, cited response.")
add_bullet(doc, "Scroll past the answer to see individual source excerpts — title, agency, year, relevance score, and the exact text.")
add_bullet(doc, "Click 'Export to Word' to download the analysis as a formatted .docx file.")

add_heading(doc, "When to use it", level=3)
add_bullet(doc, "Preparing a policy submission — find what AHURI recommends on supply-side reform, stamp duty, or planning.")
add_bullet(doc, "Writing a grant application — find evidence for unmet need in a specific demographic or geography.")
add_bullet(doc, "Briefing a board — get a 3-paragraph evidence summary on any current housing issue.")
add_bullet(doc, "Responding to a media inquiry or parliamentary question — fast, cited, authoritative.")

add_tip(doc, "The pipeline must be indexed before Ask the Research will work. If you see a warning about "
         "no reports indexed, go to Home and run the pipeline first.")


# 5.6 Policy Impact
add_heading(doc, "5.6  Policy Impact", level=2)
add_page_intro_table(doc,
    page_name   = "Policy Impact",
    tagline     = "Select any major federal housing program and get an evidence-based impact assessment.",
    who         = "Policy managers, government stakeholders, researchers assessing what prior programs delivered.",
    what_you_get= "AI impact assessment grounded in indexed research, evidence year distribution chart, exportable to Word.",
)

add_heading(doc, "How to use it", level=3)
add_bullet(doc, "Select a policy from the dropdown — all major federal housing programs are pre-loaded, including HAFF, "
           "NRAS, Nation Building, NHFIC, First Home Owner Grant, HomeBuilder, and more.")
add_bullet(doc, "Metrics for the selected policy — year, investment amount, and program type — are shown automatically.")
add_bullet(doc, "Optionally type a custom policy name (e.g., a state program or less-known initiative) in the text field.")
add_bullet(doc, "Click 'Analyse Impact'. Claude searches the index for evidence about this specific program and returns "
           "a structured impact assessment: what was intended, what was delivered, what the research says about outcomes.")
add_bullet(doc, "The evidence distribution chart shows which years of research the assessment draws from.")
add_bullet(doc, "Export to Word — formatted impact report with citations.")

add_tip(doc, "Policy Impact is particularly useful before a policy debate, a funding pitch, or a Senate Estimates hearing. "
         "It surfaces what independent researchers found — not what the government claimed.")


# 5.7 Outcome Ledger
add_heading(doc, "5.7  Outcome Ledger", level=2)
add_page_intro_table(doc,
    page_name   = "Outcome Ledger",
    tagline     = "Investment vs reality — what each major program promised and what it actually delivered.",
    who         = "CEOs, policy managers, investors, board members assessing the track record of public investment.",
    what_you_get= "Side-by-side target vs actual for every tracked program, with status flags and program detail.",
)

add_heading(doc, "What's on this page", level=3)
add_bullet(doc, "Program selector — choose any tracked housing program from the left panel.")
add_bullet(doc, "Program overview — funding amount, investment period, responsible agency, and program type.")
add_bullet(doc, "Metrics grid — every tracked commitment for the program displayed as: Metric / Target year / Target / Actual / "
           "Status (On Track, Partially Met, Missed, Exceeded, Ongoing).")
add_bullet(doc, "Research insights — any research-extracted findings about this program's real-world outcomes.")
add_bullet(doc, "Run Research Extraction — triggers a Claude analysis of indexed research for this program. "
           "Disabled if already run today (prevents repeat API usage on the same day).")

add_heading(doc, "Programs tracked", level=3)
add_bullet(doc, "Housing Australia Future Fund (HAFF) — $10B, 30,000 homes target")
add_bullet(doc, "National Rental Affordability Scheme (NRAS) — 50,000 tenancies target")
add_bullet(doc, "Nation Building Economic Stimulus Plan — social housing construction")
add_bullet(doc, "National Housing Finance and Investment Corporation (NHFIC) / Housing Australia")
add_bullet(doc, "Social Housing Accelerator — $2B state direct grants")
add_bullet(doc, "Help to Buy and other Home Guarantee Scheme programs")

add_tip(doc, "The Outcome Ledger is deliberately conservative — it tracks only what can be verified from "
         "official reporting. Where outcomes are disputed or incomplete, the status is shown as 'Partially Met' or 'Ongoing'.")


# 5.8 Policy Timeline
add_heading(doc, "5.8  Policy Timeline", level=2)
add_page_intro_table(doc,
    page_name   = "Policy Timeline",
    tagline     = "Visual chronology of every major Australian housing policy since 1990.",
    who         = "Policy researchers, advocates, government stakeholders needing historical context.",
    what_you_get= "Interactive timeline with investment amounts, program types, and links to Outcome Ledger entries.",
)

add_heading(doc, "What's on this page", level=3)
add_bullet(doc, "Horizontal timeline spanning 1990 to present — each major policy is plotted by year.")
add_bullet(doc, "Colour-coded by type — investment fund, construction program, demand subsidy, regulatory change.")
add_bullet(doc, "Click any event to see the year, investment amount, and program category.")
add_bullet(doc, "Useful for identifying patterns — gaps in investment, policy cycles, and the relationship between "
           "program eras and waitlist trends.")


# 5.9 Browse Reports
add_heading(doc, "5.9  Browse Reports", level=2)
add_page_intro_table(doc,
    page_name   = "Browse Reports",
    tagline     = "The full indexed library — filter, search, open, and download any report.",
    who         = "All users needing to find or download a specific publication.",
    what_you_get= "Filterable report table with source agency, year, type, open link, and PDF download.",
)

add_heading(doc, "What's on this page", level=3)
add_bullet(doc, "Search bar — filter by title keyword across all indexed reports.")
add_bullet(doc, "Source filter — narrow by AHURI, AIHW, ABS, Treasury, Housing Australia, Productivity Commission, DSS, or Power Housing.")
add_bullet(doc, "Year range filter — show only reports from a specific period.")
add_bullet(doc, "Report type filter — Final Report, Policy Bulletin, Research Brief, Budget Paper, Annual Report, etc.")
add_bullet(doc, "Table columns: Title, Source, Year, Type, Open (link to source URL), Download (PDF if available locally).")
add_bullet(doc, "Deduplication — each unique report URL appears only once regardless of how many times it was crawled.")
add_bullet(doc, "Dead link detection — PDF links from DSS, Treasury, and AIHW that no longer resolve are flagged.")

add_tip(doc, "PDF downloads are only available for documents downloaded during the pipeline run. "
         "For reports where only the landing page was indexed (not the PDF), use the 'Open' link to access the source website.")


# 5.10 Weekly Digest
add_heading(doc, "5.10  Weekly Digest", level=2)
add_page_intro_table(doc,
    page_name   = "Weekly Digest",
    tagline     = "An AI-generated weekly briefing — live data highlights and newly indexed research in one read.",
    who         = "CEOs, board members, communications teams, leadership teams wanting a weekly sector update.",
    what_you_get= "Formatted digest with national supply numbers, state snapshots, SHS highlights, and new research summary.",
)

add_heading(doc, "What's on this page", level=3)
add_bullet(doc, "Digest generation — click 'Generate This Week's Digest' to produce a fresh briefing based on current data.")
add_bullet(doc, "Sections: National Supply Overview, State Spotlights (all states), Homelessness & SHS Highlights, "
           "Recent Research, and a Policy Watch note.")
add_bullet(doc, "All figures sourced live from ABS and AIHW — no manual data entry required.")
add_bullet(doc, "Export to Word — the digest can be downloaded as a formatted .docx ready for email or board distribution.")
add_bullet(doc, "Date-stamped — each digest shows the generation date and data freshness note.")

add_tip(doc, "The Weekly Digest uses the Claude API. Each generation costs approximately $0.02–0.05 AUD depending on "
         "the length of the synthesis. Generate weekly, not daily.")

doc.add_page_break()


# ── 6  ROLE GUIDE ────────────────────────────────────────────────────────────

add_heading(doc, "6.  Who Uses HIVE — Role Guide", level=1)

add_para(doc, (
    "HIVE is designed for multiple roles within and around the community housing sector. "
    "The table below shows where each role should start and which pages deliver the most value."
), space_after=8)

roles = [
    {
        "title": "CEO / Executive Director",
        "org":   "CHP or peak body",
        "start": "Live Dashboard",
        "pages": "Live Dashboard → State Demand & Supply → Outcome Ledger → Weekly Digest",
        "use":   "Big-picture numbers for board meetings, investor conversations, media statements, and annual reporting.",
    },
    {
        "title": "Policy & Advocacy Manager",
        "org":   "CHP, peak body, or government",
        "start": "Ask the Research",
        "pages": "Ask the Research → Policy Impact → State Demand & Supply → Browse Reports",
        "use":   "Evidence base for submissions, policy analysis, advocacy campaigns, and parliamentary input.",
    },
    {
        "title": "Development Manager",
        "org":   "CHP — property and pipeline",
        "start": "State Demand & Supply",
        "pages": "State Demand & Supply → HAFF Investment Tracker → Outcome Ledger",
        "use":   "Demand case for new projects — waitlist demographics, bedroom mix need, and HAFF benchmarking.",
    },
    {
        "title": "Grants & Funding Officer",
        "org":   "CHP or community organisation",
        "start": "Ask the Research",
        "pages": "Ask the Research → State Demand & Supply → Browse Reports",
        "use":   "Evidence base for grant applications — unmet need statistics, demographic data, research citations.",
    },
    {
        "title": "Impact Investor / Funder",
        "org":   "Super funds, banks, philanthropies",
        "start": "Outcome Ledger",
        "pages": "Outcome Ledger → State Demand & Supply → Policy Timeline → HAFF Investment Tracker",
        "use":   "Track record of government investment, scale of unmet need, and pipeline of future demand.",
    },
    {
        "title": "Government Stakeholder",
        "org":   "State or federal housing departments",
        "start": "Policy Impact",
        "pages": "Policy Impact → State Demand & Supply → Ask the Research → Policy Timeline",
        "use":   "What prior programs delivered before designing the next one; sector recommendations from independent research.",
    },
    {
        "title": "Communications / Media",
        "org":   "CHP, peak body, or government",
        "start": "Live Dashboard",
        "pages": "Live Dashboard → Home (crisis numbers) → Ask the Research",
        "use":   "Fast, citable statistics for media releases, op-eds, and stakeholder communications.",
    },
]

tbl = doc.add_table(rows=len(roles)+1, cols=4)
tbl.style = "Table Grid"

# Header
headers = ["Role", "Where to start", "Key pages", "Primary use case"]
for cell, hdr in zip(tbl.rows[0].cells, headers):
    set_cell_bg(cell, "0F0F1A")
    cell.text = hdr
    cell.paragraphs[0].runs[0].font.color.rgb = GOLD
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)

for i, role in enumerate(roles, 1):
    row = tbl.rows[i]
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, "F5F5F5")
    row.cells[0].text = f"{role['title']}\n{role['org']}"
    row.cells[1].text = role["start"]
    row.cells[2].text = role["pages"]
    row.cells[3].text = role["use"]
    for j, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = DKGREY
        if j == 0:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()
doc.add_page_break()


# ── 7  DATA SOURCES ──────────────────────────────────────────────────────────

add_heading(doc, "7.  Data Sources", level=1)

add_para(doc, (
    "Every answer HIVE gives is grounded in real publications from the sources below. "
    "Nothing is fabricated. Every AI synthesis can be traced to a specific indexed document."
), space_after=8)

sources = [
    {
        "name":   "AHURI — Australian Housing and Urban Research Institute",
        "color":  "E74C3C",
        "desc":   "15 years of final reports, policy bulletins, research briefs, and evidence reviews. "
                  "The authoritative academic source on Australian housing. Over 500 publications indexed.",
        "update": "Ongoing — new reports crawled each pipeline run",
        "types":  "Final Reports, Policy Bulletins, Research Briefs, Evidence Reviews",
    },
    {
        "name":   "Housing Australia (formerly NHFIC)",
        "color":  "3498DB",
        "desc":   "Annual reports, Home Guarantee Scheme trends, bond aggregation data, and social housing "
                  "investment reports. Data on HAFF rounds sourced here.",
        "update": "Annual reports; HAFF data updated post each round announcement",
        "types":  "Annual Reports, State of the Nation's Housing, HAFF media releases",
    },
    {
        "name":   "Treasury — Australian Government",
        "color":  "27AE60",
        "desc":   "Federal Budget Papers (2010–2026). Budget Paper 2 lists every housing program, its funding, "
                  "and year-by-year allocations. The financial ground truth for policy accountability.",
        "update": "Annual (May Budget)",
        "types":  "Budget Papers 1–4, Mid-Year Economic and Fiscal Outlook (MYEFO)",
    },
    {
        "name":   "ABS — Australian Bureau of Statistics",
        "color":  "F39C12",
        "desc":   "Building approvals (monthly live feed), Census housing data, residential property price "
                  "indexes, housing occupancy and costs surveys.",
        "update": "Building approvals: monthly live. Census: 5-yearly. Other: annual.",
        "types":  "Cat. 8731.0 Building Approvals, Census Table Builder, RPPI",
    },
    {
        "name":   "AIHW — Australian Institute of Health and Welfare",
        "color":  "9B59B6",
        "desc":   "Specialist Homelessness Services annual reports, homelessness estimates from Census, "
                  "Indigenous housing data, housing assistance data. The authoritative source on housing outcomes.",
        "update": "SHS annual report: each December. Other: annual.",
        "types":  "SHS Annual Report, Homelessness in Australia, Housing Assistance in Australia",
    },
    {
        "name":   "Productivity Commission",
        "color":  "1ABC9C",
        "desc":   "Major housing inquiries including the landmark 2022 Housing and Homelessness report, "
                  "rental assistance review, and Report on Government Services (housing chapter).",
        "update": "Inquiry reports: as published. ROGS: annual.",
        "types":  "Inquiry Reports, Draft Reports, Report on Government Services",
    },
    {
        "name":   "DSS — Department of Social Services",
        "color":  "E67E22",
        "desc":   "National Housing and Homelessness Agreement, National Rental Affordability Scheme "
                  "documentation, homelessness strategy policy papers.",
        "update": "Policy documents: as published",
        "types":  "NHAS documentation, NRAS reports, Homelessness Strategy papers",
    },
    {
        "name":   "Power Housing Australia",
        "color":  "95A5A6",
        "desc":   "Community housing sector peak body publications and State of the Sector reports "
                  "where accessible. Supplements the government and academic sources with sector voice.",
        "update": "Annual",
        "types":  "State of the Sector, Sector Capacity reports",
    },
]

for src in sources:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "F9F9F9")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), src["color"])
    tcBdr.append(left)
    tcPr.append(tcBdr)

    cell.paragraphs[0].clear()

    def _row(text, bold=False, clr=DKGREY, sz=10):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(sz)
        r.font.color.rgb = clr

    _row(src["name"], bold=True, clr=NAVY, sz=11)
    _row(src["desc"])
    _row(f"Update frequency: {src['update']}", clr=RGBColor(0x77, 0x77, 0x77))
    _row(f"Report types: {src['types']}", clr=RGBColor(0x77, 0x77, 0x77))

    doc.add_paragraph()

doc.add_page_break()


# ── 8  TIPS & FAQ ────────────────────────────────────────────────────────────

add_heading(doc, "8.  Tips & Frequently Asked Questions", level=1)

faqs = [
    (
        "How often does the live data update?",
        "ABS building approvals refresh monthly — the platform fetches fresh data when the existing "
        "cache is more than 30 days old. AIHW SHS data is annual (the latest is 2023–24). "
        "State waitlist data is sourced from state authority annual reports — also annual. "
        "To force a refresh, delete the relevant .json file from data/live_cache/ and reload the page."
    ),
    (
        "The AI answer doesn't seem right — what should I do?",
        "Check the source excerpts below the answer — each chunk includes the original text, "
        "publication title, agency, and year. If the answer seems off, it's usually because the most "
        "relevant publication wasn't indexed, or the question was too broad. Try narrowing the year "
        "range in Search Filters or rephrasing your question more specifically."
    ),
    (
        "How do I export results to Word?",
        "After generating any AI synthesis (Ask the Research, Policy Impact, or Weekly Digest), "
        "a blue 'Export to Word (.docx)' button appears below the answer. Click it to download "
        "a formatted document with the synthesis, sources, and metadata."
    ),
    (
        "Can I add new reports to the index?",
        "Yes. Place PDF or HTML files in the data/reports/ directory and re-run the pipeline from "
        "the Home page. HIVE will detect new files and index them. For web sources, the crawler "
        "configuration is in crawler/run_all.py."
    ),
    (
        "Does HIVE send my data anywhere?",
        "Only when you request an AI synthesis or insight. The text of relevant document chunks is "
        "sent to the Claude API (Anthropic). No personal data, report files, or waitlist data leave "
        "your machine. All other features — charts, tables, live data — run entirely locally."
    ),
    (
        "The pipeline is taking very long — is that normal?",
        "Yes on first run. Crawling and indexing 681+ reports takes 15–30 minutes depending on "
        "internet speed and machine performance. The pipeline can be left running in the background. "
        "Subsequent runs are much faster as previously indexed documents are skipped."
    ),
    (
        "Why is the Run Pipeline button not appearing?",
        "The run button only appears when the vector index is empty (0 chunks). If you have already "
        "run the pipeline, HIVE treats the index as current and doesn't offer a re-run — this prevents "
        "unnecessary API costs. To force a re-index, delete the data/chroma/ directory and reload."
    ),
    (
        "Some 'Browse Reports' links return 404 errors — why?",
        "Government departments (DSS, Treasury, AIHW) periodically reorganise their websites and "
        "change PDF URLs. HIVE flags known dead links in the Browse Reports table. The document "
        "content is still searchable via Ask the Research because it was indexed before the link changed."
    ),
    (
        "Can I use HIVE for external publications — e.g., board papers or investor decks?",
        "Yes, with appropriate attribution. All data originates from public government and research "
        "publications. When citing figures from HIVE, reference the underlying source (e.g., "
        "'ABS Building Approvals Cat. 8731.0, March 2025') rather than HIVE itself."
    ),
    (
        "How is the 'Years to clear waitlist' figure calculated?",
        "It divides the current approved applicant count by the average annual social housing "
        "completions for that state over the last 3 years. It is an indicative planning metric, "
        "not a forecast — it does not account for new applicants joining the list each year."
    ),
]

for q, a in faqs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"Q: {q}")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(f"A: {a}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = DKGREY

doc.add_page_break()


# ── 9  GLOSSARY ──────────────────────────────────────────────────────────────

add_heading(doc, "9.  Glossary", level=1)

terms = [
    ("ABS",              "Australian Bureau of Statistics — produces building approvals, Census, and price data."),
    ("AHURI",            "Australian Housing and Urban Research Institute — the primary academic research body "
                          "for Australian housing. Funded jointly by the federal government and state governments."),
    ("AIHW",             "Australian Institute of Health and Welfare — produces the SHS annual report and "
                          "homelessness statistics."),
    ("Approved applicant","A person or household registered on a state social housing waitlist who has been "
                          "assessed as eligible and is awaiting allocation."),
    ("CHP",              "Community Housing Provider — a registered not-for-profit organisation that provides "
                          "social and affordable housing."),
    ("DSS",              "Department of Social Services — federal department responsible for housing and "
                          "homelessness policy and the NHAS."),
    ("HAFF",             "Housing Australia Future Fund — a $10 billion off-budget investment fund established "
                          "in 2023, targeting 30,000 social and affordable homes over 5 years."),
    ("National Accord",  "National Housing Accord — a September 2022 agreement between federal, state, and "
                          "territory governments targeting 1.2 million new homes over 5 years (240,000/year)."),
    ("NHFIC",            "National Housing Finance and Investment Corporation — renamed Housing Australia in 2023. "
                          "Provides bond aggregation lending and administers the Home Guarantee Scheme and HAFF."),
    ("NHAS",             "National Housing and Homelessness Agreement — the funding and policy framework between "
                          "the Commonwealth and states for housing assistance."),
    ("NRAS",             "National Rental Affordability Scheme — a 2008 program targeting 50,000 affordable rental "
                          "tenancies through tax incentives to investors. Wound down from 2014."),
    ("Run rate",         "In the context of HIVE, the annualised dwelling approval rate: the 3-month average of "
                          "ABS monthly building approvals multiplied by 12."),
    ("SHS",              "Specialist Homelessness Services — crisis and transitional housing services funded under "
                          "the NHAS. The AIHW publishes annual SHS client outcome data."),
    ("Social housing",   "Government-owned or community housing provided at below-market rent to low-income "
                          "households. Distinct from affordable housing, which is typically at 75–80% of market rent."),
    ("Vector index",     "The database HIVE uses to store AI-encoded document chunks. When you search, HIVE finds "
                          "chunks that are semantically similar to your question — not just keyword matches."),
    ("Waitlist",         "Colloquial term for the social housing register of approved applicants awaiting allocation. "
                          "The formal term varies by state (e.g., 'Housing Register' in NSW and WA)."),
]

tbl = doc.add_table(rows=len(terms)+1, cols=2)
tbl.style = "Table Grid"

hdr = tbl.rows[0].cells
for cell, txt in zip(hdr, ["Term", "Definition"]):
    set_cell_bg(cell, "0F0F1A")
    cell.text = txt
    cell.paragraphs[0].runs[0].font.color.rgb = GOLD
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

for i, (term, defn) in enumerate(terms, 1):
    row = tbl.rows[i]
    if i % 2 == 0:
        set_cell_bg(row.cells[0], "F5F5F5")
        set_cell_bg(row.cells[1], "F5F5F5")
    row.cells[0].text = term
    row.cells[1].text = defn
    row.cells[0].paragraphs[0].runs[0].bold = True
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = DKGREY


# ── FOOTER NOTE ──────────────────────────────────────────────────────────────

doc.add_page_break()

add_para(doc, "", space_before=40)
add_para(doc, "HIVE — Housing Intelligence & Evidence", bold=True, size=12,
         color=NAVY, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, f"User Guide v1.0  ·  {TODAY}",
         size=9, color=DKGREY, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Prepared by Sunny Kim  ·  Housing Data Lead  ·  Community Housing Professional, Australia",
         size=9, color=DKGREY, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "linkedin.com/in/sunny-kim-58a780100",
         size=9, color=DKGREY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_para(doc,
    "This document is intended for internal use by housing sector professionals. "
    "All data sourced from publicly available government and research publications. "
    "AI-generated content should be verified against primary sources before formal submission use.",
    size=8, color=RGBColor(0xAA, 0xAA, 0xAA),
    alignment=WD_ALIGN_PARAGRAPH.CENTER)


# ── Save ─────────────────────────────────────────────────────────────────────

output_path = "HIVE_User_Guide.docx"
doc.save(output_path)
print(f"✅  Saved: {output_path}")
